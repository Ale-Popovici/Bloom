import os
import uuid
from fastapi import UploadFile
import fitz  # PyMuPDF
import docx
import tempfile
import logging
import re
from typing import Dict, Any, Optional

from services.vector_store import get_collection, add_documents
from utils.text_splitter import split_text

# Set up logging with more detail
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Track processing status
processing_status = {}


async def process_document(file: UploadFile, module_code: Optional[str] = None) -> str:
    """
    Process a document and add it to the vector store.
    This is a simplified version that doesn't use a background queue.

    Args:
        file (UploadFile): The uploaded file
        module_code (str, optional): Module code for collection organization

    Returns:
        str: The document ID
    """
    # Generate unique ID for the document
    document_id = str(uuid.uuid4())

    # Determine collection name
    collection_name = f"module_{module_code}" if module_code else "bloom_documents"

    logger.info(
        f"Processing document '{file.filename}' with ID {document_id} for collection '{collection_name}'")

    # Update processing status
    processing_status[document_id] = {
        "filename": file.filename,
        "status": "processing",
        "progress": 0,
        "collection": collection_name
    }

    try:
        # Extract text based on file type
        if file.filename.lower().endswith('.pdf'):
            text = await extract_text_from_pdf(file)
            logger.info(
                f"Extracted {len(text)} characters from PDF '{file.filename}'")
        elif file.filename.lower().endswith(('.docx', '.doc')):
            text = await extract_text_from_docx(file)
            logger.info(
                f"Extracted {len(text)} characters from DOCX/DOC '{file.filename}'")
        else:
            raise ValueError(f"Unsupported file format: {file.filename}")

        # Check for empty extraction result
        # Minimum threshold for useful content
        if not text or len(text.strip()) < 100:
            logger.warning(
                f"Extracted text too short or empty for '{file.filename}'. Length: {len(text)}")
            # Try alternative extraction if first method failed
            if file.filename.lower().endswith('.pdf'):
                logger.info(
                    f"Attempting alternative PDF extraction method for '{file.filename}'")
                await file.seek(0)  # Reset file pointer
                text = await extract_pdf_with_alternative_method(file)
                logger.info(
                    f"Alternative extraction yielded {len(text)} characters")
            elif file.filename.lower().endswith(('.docx', '.doc')):
                logger.info(
                    f"Attempting alternative DOCX/DOC extraction method for '{file.filename}'")
                await file.seek(0)  # Reset file pointer
                text = await extract_docx_with_alternative_method(file)
                logger.info(
                    f"Alternative extraction yielded {len(text)} characters")

            # Final check for minimum text content
            if not text or len(text.strip()) < 50:
                logger.error(
                    f"Failed to extract meaningful text from '{file.filename}'")
                processing_status[document_id]["status"] = "warning"
                processing_status[document_id]["warning"] = "Limited text extraction"
                # Use filename and metadata as minimum content
                text = f"Document: {file.filename}\nType: {file.filename.split('.')[-1].upper()}\n"
                text += f"Module: {module_code if module_code else 'General'}\n"
                text += "Content extraction limited - may be a scanned document or contain primarily non-text content."

        # Update progress
        processing_status[document_id]["progress"] = 30

        # Split text into chunks
        chunks = split_text(text)
        logger.info(
            f"Split text into {len(chunks)} chunks for document '{file.filename}'")

        # Update progress
        processing_status[document_id]["progress"] = 50

        # Prepare data for Chroma DB
        texts = []
        metadatas = []

        for i, chunk in enumerate(chunks):
            # Skip empty chunks
            if not chunk or len(chunk.strip()) < 20:
                logger.warning(
                    f"Skipping empty or very short chunk {i} in '{file.filename}'")
                continue

            texts.append(chunk)

            # Create metadata without None values
            metadata = {
                "document_id": document_id,
                "filename": file.filename,
                "chunk_index": i,
                "total_chunks": len(chunks)
            }

            # Only add module_code if it's not None
            if module_code is not None:
                metadata["module_code"] = module_code

            metadatas.append(metadata)

        # Check if we have any valid chunks
        if not texts:
            logger.warning(
                f"No valid text chunks extracted from '{file.filename}'")
            # Add at least one chunk with file metadata
            texts = [f"Document: {file.filename}\nType: {file.filename.split('.')[-1].upper()}\n" +
                     f"Module: {module_code if module_code else 'General'}\n" +
                     "This document may be a scanned document or contain primarily non-text content."]
            metadatas = [{
                "document_id": document_id,
                "filename": file.filename,
                "chunk_index": 0,
                "total_chunks": 1
            }]
            if module_code is not None:
                metadatas[0]["module_code"] = module_code

        # Update progress
        processing_status[document_id]["progress"] = 70

        # Add documents to ChromaDB using the specified collection
        try:
            # Print verification of what's being added
            logger.info(f"Adding {len(texts)} chunks to '{collection_name}'")
            logger.info(f"First chunk preview: {texts[0][:100]}...")
            logger.info(f"Metadata: {metadatas[0]}")

            add_documents(texts, metadatas, collection_name)
            logger.info(
                f"Successfully added {len(texts)} chunks to collection '{collection_name}' for document '{file.filename}'")
        except Exception as e:
            logger.error(
                f"Failed to add chunks to collection '{collection_name}': {str(e)}")
            raise e

        # Update status to complete
        processing_status[document_id]["status"] = "complete"
        processing_status[document_id]["progress"] = 100
        logger.info(
            f"Completed processing document '{file.filename}' with ID {document_id}")

        return document_id

    except Exception as e:
        # Update status to failed
        processing_status[document_id]["status"] = "failed"
        processing_status[document_id]["error"] = str(e)
        logger.error(f"Error processing document '{file.filename}': {str(e)}")
        raise e


async def extract_text_from_pdf(file: UploadFile) -> str:
    """
    Extract text from a PDF file

    Args:
        file (UploadFile): The PDF file

    Returns:
        str: Extracted text
    """
    # Save upload to temp file
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_path = temp_file.name

        logger.info(f"Saved PDF to temporary file: {temp_path}")

        # Extract text using PyMuPDF
        doc = fitz.open(temp_path)
        logger.info(f"Opened PDF with {len(doc)} pages")

        # More comprehensive extraction with metadata
        text_parts = []

        # Add document metadata if available
        metadata = doc.metadata
        if metadata:
            meta_text = "Document Metadata:\n"
            for key, value in metadata.items():
                if value:
                    meta_text += f"{key}: {value}\n"
            text_parts.append(meta_text)

        # Extract text from each page with page numbers
        for i, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"Page {i+1}:\n{page_text}")

        # Extract table of contents if available
        toc = doc.get_toc()
        if toc:
            toc_text = "Table of Contents:\n"
            for level, title, page in toc:
                toc_text += f"{'  ' * (level-1)}- {title} (Page {page})\n"
            text_parts.append(toc_text)

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from PDF")
        return full_text

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise e

    finally:
        # Clean up temp file
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
            logger.info(f"Removed temporary PDF file: {temp_path}")

        # Reset file pointer for potential reuse
        await file.seek(0)


async def extract_text_from_docx(file: UploadFile) -> str:
    """
    Extract text from a DOCX file

    Args:
        file (UploadFile): The DOCX file

    Returns:
        str: Extracted text
    """
    # Save upload to temp file
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_path = temp_file.name

        logger.info(f"Saved DOCX to temporary file: {temp_path}")

        # Extract text using python-docx
        doc = docx.Document(temp_path)
        logger.info(
            f"Opened DOCX with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")

        # Extract document properties
        core_properties = doc.core_properties
        text_parts = []

        # Add document metadata if available
        if core_properties:
            meta_text = "Document Metadata:\n"
            if core_properties.title:
                meta_text += f"Title: {core_properties.title}\n"
            if core_properties.author:
                meta_text += f"Author: {core_properties.author}\n"
            if core_properties.created:
                meta_text += f"Created: {core_properties.created}\n"
            if core_properties.modified:
                meta_text += f"Modified: {core_properties.modified}\n"
            if core_properties.subject:
                meta_text += f"Subject: {core_properties.subject}\n"
            if core_properties.keywords:
                meta_text += f"Keywords: {core_properties.keywords}\n"

            text_parts.append(meta_text)

        # Extract headings and paragraphs with structure
        current_heading = "Document Content"
        paragraphs_text = []

        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                if paragraphs_text:
                    text_parts.append(
                        f"{current_heading}:\n" + "\n".join(paragraphs_text))
                    paragraphs_text = []
                current_heading = para.text
            elif para.text.strip():
                paragraphs_text.append(para.text)

        # Add the last section
        if paragraphs_text:
            text_parts.append(f"{current_heading}:\n" +
                              "\n".join(paragraphs_text))

        # Extract tables
        for i, table in enumerate(doc.tables):
            table_text = f"Table {i+1}:\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                table_text += " | ".join(row_text) + "\n"
            text_parts.append(table_text)

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        return full_text

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise e

    finally:
        # Clean up temp file
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
            logger.info(f"Removed temporary DOCX file: {temp_path}")

        # Reset file pointer for potential reuse
        await file.seek(0)


async def extract_pdf_with_alternative_method(file: UploadFile) -> str:
    """
    Alternative method to extract text from PDFs that may be scanned or image-based
    """
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_path = temp_file.name

        logger.info(f"Using alternative extraction for PDF: {temp_path}")

        # First attempt: Try to get any text available in the PDF
        doc = fitz.open(temp_path)
        text_parts = []

        # Add file metadata
        metadata = doc.metadata
        if metadata:
            meta_text = "Document Metadata:\n"
            for key, value in metadata.items():
                if value:
                    meta_text += f"{key}: {value}\n"
            text_parts.append(meta_text)

        # Extract whatever text we can find, page by page
        for i, page in enumerate(doc):
            # Try get_text with different parameters
            text = page.get_text("text")  # Plain text
            if not text.strip():
                text = page.get_text("blocks")  # Try blocks mode
            if not text.strip():
                text = page.get_text("rawdict")  # Try raw dict mode
                if isinstance(text, dict) and "blocks" in text:
                    # Extract text from raw dict
                    block_texts = []
                    for block in text["blocks"]:
                        if "lines" in block:
                            for line in block["lines"]:
                                if "spans" in line:
                                    for span in line["spans"]:
                                        if "text" in span and span["text"].strip():
                                            block_texts.append(span["text"])
                    text = " ".join(block_texts)

            if text and isinstance(text, str) and text.strip():
                text_parts.append(f"Page {i+1}:\n{text}")
            else:
                # If no text extracted, note it's likely an image
                text_parts.append(
                    f"Page {i+1}: [This page appears to contain only images or non-extractable content]")

        full_text = "\n\n".join(text_parts)

        # If still no useful text, add basic document info
        if not full_text.strip() or len(full_text.strip()) < 50:
            full_text = f"Document appears to be a scanned PDF or contains primarily images.\n"
            if metadata:
                for key, value in metadata.items():
                    if value:
                        full_text += f"{key}: {value}\n"

        return full_text

    except Exception as e:
        logger.error(f"Error in alternative PDF extraction: {str(e)}")
        return f"PDF extraction failed. Error: {str(e)}"

    finally:
        # Clean up temp file
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
            logger.info(f"Removed temporary PDF file: {temp_path}")

        # Reset file pointer for potential reuse
        await file.seek(0)


async def extract_docx_with_alternative_method(file: UploadFile) -> str:
    """
    Alternative method to extract text from DOCX/DOC files that may have complex formatting
    """
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_path = temp_file.name

        logger.info(f"Using alternative extraction for DOCX: {temp_path}")

        # Try to extract using python-docx
        try:
            doc = docx.Document(temp_path)
            text_parts = []

            # Try to get document properties
            try:
                core_properties = doc.core_properties
                meta_text = "Document Metadata:\n"
                for prop in ['title', 'author', 'created', 'modified', 'subject', 'keywords']:
                    value = getattr(core_properties, prop, None)
                    if value:
                        meta_text += f"{prop.capitalize()}: {value}\n"
                text_parts.append(meta_text)
            except Exception as e:
                logger.warning(
                    f"Could not extract document properties: {str(e)}")
                text_parts.append("Document Metadata: Could not be extracted")

            # Get all text, even if not structured
            all_text = "\n".join(
                [para.text for para in doc.paragraphs if para.text.strip()])
            if all_text:
                text_parts.append("Document Content:\n" + all_text)

            # Get table content
            table_texts = []
            for i, table in enumerate(doc.tables):
                table_text = f"Table {i+1}:\n"
                for row in table.rows:
                    row_text = " | ".join(
                        [cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        table_text += row_text + "\n"
                if len(table_text) > 10:  # Only add if there's meaningful content
                    table_texts.append(table_text)

            if table_texts:
                text_parts.append("\n".join(table_texts))

            full_text = "\n\n".join(
                [part for part in text_parts if part.strip()])

            # If still no useful text, add basic document info
            if not full_text.strip() or len(full_text.strip()) < 50:
                full_text = f"Document appears to contain limited extractable text content.\n"
                try:
                    if core_properties:
                        for prop in ['title', 'author', 'created', 'modified']:
                            value = getattr(core_properties, prop, None)
                            if value:
                                full_text += f"{prop.capitalize()}: {value}\n"
                except:
                    pass

            return full_text

        except Exception as docx_error:
            logger.warning(f"python-docx extraction failed: {str(docx_error)}")

            # Try using a simpler text extraction as fallback
            try:
                with open(temp_path, 'rb') as doc_file:
                    doc_bytes = doc_file.read()

                # Try to find plain text in the binary content
                # This is a crude method but might extract some text from DOC files
                readable_text = re.findall(
                    b'[a-zA-Z0-9 .,;:\'"\-_\n\r\t]{4,}', doc_bytes)
                text = b'\n'.join(readable_text).decode(
                    'utf-8', errors='ignore')

                if text and len(text) > 50:
                    return f"Document Content (limited extraction):\n{text}"
                else:
                    return "Could not extract meaningful text from this document. It may be in a protected format or contain primarily non-text content."
            except Exception as bin_error:
                logger.error(f"Binary extraction failed: {str(bin_error)}")
                return "Document text extraction failed. The file may be corrupted or in an unsupported format."

    except Exception as e:
        logger.error(f"Error in alternative DOCX extraction: {str(e)}")
        return f"DOCX extraction failed. Error: {str(e)}"

    finally:
        # Clean up temp file
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)
            logger.info(f"Removed temporary DOCX file: {temp_path}")

        # Reset file pointer for potential reuse
        await file.seek(0)


def get_processing_status(document_id: str) -> Dict[str, Any]:
    """
    Get the processing status for a document

    Args:
        document_id (str): The document ID

    Returns:
        Dict: Status information
    """
    if document_id in processing_status:
        return processing_status[document_id]
    return {"status": "not_found"}
